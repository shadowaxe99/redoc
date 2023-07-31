import PyPDF2
from docx import Document

def parseDocument(file):
    file_extension = file.filename.rsplit('.', 1)[1].lower()
    
    if file_extension == 'txt':
        try:
            with open(file, 'r') as f:
                content = f.read()
            return content
        except Exception as e:
            print(f"Error reading file: {str(e)}")
            return None
    elif file_extension == 'pdf':
        try:
            pdf_file_obj = open(file, 'rb')
            pdf_reader = PyPDF2.PdfFileReader(pdf_file_obj)
            content = ''
            for page_num in range(pdf_reader.numPages):
                page_obj = pdf_reader.getPage(page_num)
                content += page_obj.extractText()
            pdf_file_obj.close()
            return content
        except Exception as e:
            print(f"Error reading PDF file: {str(e)}")
            return None
    elif file_extension == 'docx':
        try:
            doc = Document(file)
            content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
            return content
        except Exception as e:
            print(f"Error reading Word file: {str(e)}")
            return None
    else:
        print(f"Unsupported file extension: {file_extension}")
        return None