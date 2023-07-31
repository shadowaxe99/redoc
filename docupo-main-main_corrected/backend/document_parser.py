import PyPDF2
from docx import Document

DOCUMENT_FORMATS = ['pdf', 'docx', 'txt']

def parseDocument(file_path):
    file_extension = file_path.split('.')[-1]

    if file_extension not in DOCUMENT_FORMATS:
        raise ValueError(f"Unsupported file format. Supported formats are {DOCUMENT_FORMATS}")

    if file_extension == 'pdf':
        return parsePDF(file_path)
    elif file_extension == 'docx':
        return parseDOCX(file_path)
    elif file_extension == 'txt':
        return parseTXT(file_path)

def parsePDF(file_path):
    pdf_file_obj = open(file_path, 'rb')
    pdf_reader = PyPDF2.PdfFileReader(pdf_file_obj)
    num_pages = pdf_reader.numPages
    document_content = ''

    for page in range(num_pages):
        page_obj = pdf_reader.getPage(page)
        document_content += page_obj.extractText()

    pdf_file_obj.close()
    return document_content

def parseDOCX(file_path):
    doc = Document(file_path)
    document_content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
    return document_content

def parseTXT(file_path):
    with open(file_path, 'r') as file:
        document_content = file.read()
    return document_content
```